from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from src.docgen.generator import GeneratedDocument
from src.docgen.schemas import get_template_schemas


def _font() -> str:
    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont("JarvisUnicode", str(path)))
                return "JarvisUnicode"
            except Exception:
                continue
    return "Helvetica"


def export_docx(generated: GeneratedDocument) -> bytes:
    schema = get_template_schemas()[generated.template_type]
    document = Document()
    title = document.add_heading(schema.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning = document.add_paragraph("ДЕМОНСТРАЦИОННЫЙ ЧЕРНОВИК")
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details = document.add_table(rows=0, cols=2)
    details.style = "Table Grid"
    for spec in schema.fields:
        value = generated.fields.get(spec.name, spec.default)
        if value in (None, ""):
            continue
        cells = details.add_row().cells
        cells[0].text = spec.label
        cells[1].text = str(value)
    if generated.items:
        document.add_heading("Позиции", level=1)
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["№", "Наименование", "Артикул", "Кол-во", "Цена", "Сумма"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in generated.items:
            cells = table.add_row().cells
            values = [
                item.get("no", ""),
                item.get("наименование", ""),
                item.get("артикул", ""),
                item.get("количество", ""),
                f"{item.get('цена', 0):.2f}",
                f"{item.get('сумма', 0):.2f}",
            ]
            for index, value in enumerate(values):
                cells[index].text = str(value)
        document.add_paragraph(f"Итого: {generated.total:.2f} руб.")
    document.add_paragraph("Создано на синтетических данных. Требуется проверка юристом.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def export_pdf(generated: GeneratedDocument) -> bytes:
    schema = get_template_schemas()[generated.template_type]
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    _, height = A4
    font_name = _font()
    pdf.setTitle(schema.title)
    pdf.setFont(font_name, 16)
    pdf.drawString(48, height - 60, schema.title)
    pdf.setFont(font_name, 9)
    pdf.drawString(48, height - 80, "ДЕМОНСТРАЦИОННЫЙ ЧЕРНОВИК")
    y = height - 110
    for spec in schema.fields:
        value = generated.fields.get(spec.name, spec.default)
        if value in (None, ""):
            continue
        line = f"{spec.label}: {value}".replace("\n", " ")
        if font_name == "Helvetica":
            line = line.encode("latin-1", "replace").decode("latin-1")
        for start in range(0, len(line), 92):
            if y < 55:
                pdf.showPage()
                pdf.setFont(font_name, 9)
                y = height - 55
            pdf.drawString(48, y, line[start : start + 92])
            y -= 14
    if generated.items:
        y -= 8
        pdf.setFont(font_name, 11)
        pdf.drawString(48, y, "Позиции")
        y -= 18
        pdf.setFont(font_name, 9)
        for item in generated.items:
            line = (
                f"{item.get('no')}. {item.get('наименование')} — "
                f"{item.get('количество')} x {item.get('цена'):.2f} = {item.get('сумма'):.2f}"
            )
            if y < 55:
                pdf.showPage()
                pdf.setFont(font_name, 9)
                y = height - 55
            pdf.drawString(48, y, line[:100])
            y -= 14
        pdf.drawString(48, y - 6, f"Итого: {generated.total:.2f} руб.")
    pdf.setFont(font_name, 8)
    pdf.drawString(48, 30, "Синтетические данные. Требуется проверка юристом.")
    pdf.save()
    return output.getvalue()
