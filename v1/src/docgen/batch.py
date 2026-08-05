from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    original = paragraph.text
    updated = original
    for source, target in replacements.items():
        updated = updated.replace(source, target)
    if updated == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def replace_document_text(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)


def generate_batch(template_bytes: bytes, jobs: list[dict]) -> bytes:
    if not jobs:
        raise ValueError("At least one batch job is required")
    archive_buffer = BytesIO()
    with ZipFile(archive_buffer, "w", compression=ZIP_DEFLATED) as archive:
        for index, job in enumerate(jobs, 1):
            replacements = {str(key): str(value) for key, value in dict(job.get("replace", {})).items()}
            document = Document(BytesIO(template_bytes))
            replace_document_text(document, replacements)
            output = BytesIO()
            document.save(output)
            raw_name = str(job.get("output_name") or f"document_{index}")
            safe_name = "".join(char if char.isalnum() or char in "-_ " else "_" for char in raw_name).strip()
            archive.writestr(f"{safe_name or f'document_{index}'}.docx", output.getvalue())
    return archive_buffer.getvalue()
