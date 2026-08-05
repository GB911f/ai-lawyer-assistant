from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from src.docgen.batch import generate_batch
from src.docgen.exporters import export_docx, export_pdf
from src.docgen.generator import generate_document
from src.docgen.schemas import schemas_as_dict

GOLDEN_DIR = Path(__file__).parent / "golden"


def test_schemas_contain_only_public_demo_identity() -> None:
    serialized = str(schemas_as_dict()).lower()
    assert "ооо «север»" in serialized
    assert "example.test" in serialized


def test_docx_and_pdf_exports_are_valid_binary_documents() -> None:
    generated = generate_document(
        "kommercheskoe_predlozhenie",
        {"название_организации": "ООО «Север»", "номер_документа": "DEMO-1"},
        [{"наименование": "Комплект", "количество": "2", "цена": "100"}],
    )
    docx = export_docx(generated)
    pdf = export_pdf(generated)
    assert docx.startswith(b"PK")
    assert pdf.startswith(b"%PDF")
    parsed = Document(BytesIO(docx))
    assert "Коммерческое предложение" in "\n".join(paragraph.text for paragraph in parsed.paragraphs)


def test_batch_replaces_placeholders() -> None:
    template = Document()
    template.add_paragraph("Заказчик: {{customer_name}}")
    source = BytesIO()
    template.save(source)
    archive = generate_batch(
        source.getvalue(),
        [{"output_name": "contract_demo", "replace": {"{{customer_name}}": "ООО «Партнёр»"}}],
    )
    with ZipFile(BytesIO(archive)) as zip_file:
        names = zip_file.namelist()
        assert names == ["contract_demo.docx"]
        document = Document(BytesIO(zip_file.read(names[0])))
    assert "ООО «Партнёр»" in document.paragraphs[0].text


def test_generated_markdown_matches_golden_file() -> None:
    generated = generate_document(
        "kommercheskoe_predlozhenie",
        {"название_организации": "ООО «Север»", "номер_документа": "DEMO-GOLDEN"},
        [{"наименование": "Комплект", "количество": "2", "цена": "100"}],
    )
    expected = (GOLDEN_DIR / "kommercheskoe_predlozhenie.txt").read_text(encoding="utf-8")
    assert generated.markdown == expected.rstrip("\n")
