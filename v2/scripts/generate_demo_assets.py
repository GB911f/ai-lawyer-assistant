from pathlib import Path

from docx import Document


def main() -> None:
    output = Path("data/demo/templates/contract_template.docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("ДЕМОНСТРАЦИОННЫЙ ДОГОВОР № {{contract_number}}", level=0)
    document.add_paragraph("Поставщик: {{supplier_name}}, ИНН {{supplier_inn}}")
    document.add_paragraph("Заказчик: {{customer_name}}, ИНН {{customer_inn}}")
    document.add_paragraph("Предмет: {{subject}}")
    document.add_paragraph("Срок поставки: {{delivery_term}}")
    document.add_paragraph("Черновик на синтетических данных. Требуется проверка юристом.")
    document.save(output)
    print(f"Created {output}")


if __name__ == "__main__":
    main()
