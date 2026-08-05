from pathlib import Path

from docx import Document

from src.indexer.parsers import parse_file


def test_text_parser(tmp_path: Path) -> None:
    path = tmp_path / "document.txt"
    path.write_text("Демонстрационный договор", encoding="utf-8")
    result = parse_file(path)
    assert result.parser == "text"
    assert result.text == "Демонстрационный договор"


def test_docx_parser_reads_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "document.docx"
    document = Document()
    document.add_paragraph("Договор поставки")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Срок"
    table.cell(0, 1).text = "20 дней"
    document.save(path)
    result = parse_file(path)
    assert result.parser == "docx"
    assert "Договор поставки" in result.text
    assert "Срок | 20 дней" in result.text
