from pathlib import Path

import httpx
from docx import Document

from src.indexer.ocr.qwen_vl import QwenVLOCR
from src.indexer.parsers import parse_file


def test_text_parser(tmp_path: Path) -> None:
    path = tmp_path / "document.txt"
    path.write_text("Демонстрационный договор", encoding="utf-8")
    result = parse_file(path)
    assert result.parser == "text"
    assert result.text == "Демонстрационный договор"


def test_docx_parser_reads_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "document.docx"
    document = Document()
    document.add_paragraph("Договор поставки")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Срок"
    table.cell(0, 1).text = "20 дней"
    document.save(path)
    result = parse_file(path)
    assert result.parser == "docx"
    assert "Договор поставки" in result.text
    assert "Срок | 20 дней" in result.text


def test_qwen_vl_ocr_calls_local_ollama(tmp_path: Path) -> None:
    image = tmp_path / "scan.png"
    image.write_bytes(b"synthetic-image-bytes")

    def handler(request: httpx.Request) -> httpx.Response:
        assert request.url.path == "/api/generate"
        body = __import__("json").loads(request.content)
        assert body["model"] == "qwen3-vl:4b"
        assert body["images"]
        return httpx.Response(200, json={"response": "ДОГОВОР № DEMO-1\nИНН 7700000000"})

    client = httpx.Client(transport=httpx.MockTransport(handler))
    result = QwenVLOCR(client=client, attempts=1).extract(image)
    client.close()
    assert result.engine == "qwen_vl"
    assert "ДОГОВОР" in result.text
