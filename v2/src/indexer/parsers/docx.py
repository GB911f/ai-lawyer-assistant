from pathlib import Path

from docx import Document


def parse_docx(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(chunk for chunk in chunks if chunk).strip()
